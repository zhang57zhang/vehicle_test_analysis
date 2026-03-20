"""
Word document report generator.
"""

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd


@dataclass
class ReportSection:
    """A section in the report."""

    title: str
    content: str
    level: int = 1  # 1 = heading 1, 2 = heading 2, etc.


@dataclass
class ReportData:
    """Data container for report generation."""

    title: str
    project_name: str
    test_phase: str
    generated_at: datetime
    generated_by: str
    sections: List[ReportSection]
    tables: List[Dict[str, Any]]
    figures: List[Dict[str, Any]]
    metadata: Dict[str, Any]


class WordReportGenerator:
    """
    Generator for Word (.docx) format reports.
    """

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize Word report generator.

        Args:
            template_path: Optional path to Word template file.
        """
        self.template_path = template_path
        self._document = None

    def generate(
        self,
        report_data: ReportData,
        output_path: Path,
        include_data_source: bool = True,
    ) -> bool:
        """
        Generate Word report.

        Args:
            report_data: Data for the report.
            output_path: Path to save the report.
            include_data_source: Whether to include data source information.

        Returns:
            True if generation successful.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for Word report generation. "
                "Install with: pip install python-docx"
            )

        # Create document
        if self.template_path and self.template_path.exists():
            doc = Document(str(self.template_path))
        else:
            doc = Document()

        # Add title
        title = doc.add_heading(report_data.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        doc.add_paragraph(f"项目名称: {report_data.project_name}")
        doc.add_paragraph(f"测试阶段: {report_data.test_phase}")
        doc.add_paragraph(f"生成时间: {report_data.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"生成人: {report_data.generated_by}")

        doc.add_page_break()

        # Add sections
        for section in report_data.sections:
            doc.add_heading(section.title, level=section.level)
            doc.add_paragraph(section.content)

        # Add tables
        for table_data in report_data.tables:
            self._add_table(doc, table_data)

        # Add figures
        for figure_data in report_data.figures:
            self._add_figure(doc, figure_data)

        # Add data source information
        if include_data_source and report_data.metadata:
            doc.add_page_break()
            doc.add_heading("数据溯源", level=1)
            for key, value in report_data.metadata.items():
                doc.add_paragraph(f"{key}: {value}")

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return True

    def _add_table(self, doc, table_data: Dict[str, Any]) -> None:
        """Add a table to the document."""
        from docx.shared import Pt

        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        title = table_data.get("title", "")

        if title:
            doc.add_paragraph(title, style="Heading 3")

        if not headers or not rows:
            return

        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"

        # Add headers
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                table.rows[row_idx + 1].cells[col_idx].text = str(cell_data)

    def _add_figure(self, doc, figure_data: Dict[str, Any]) -> None:
        """Add a figure to the document."""
        from docx.shared import Inches

        image_path = figure_data.get("path")
        caption = figure_data.get("caption", "")
        width = figure_data.get("width", 6)

        if image_path and Path(image_path).exists():
            doc.add_picture(str(image_path), width=Inches(width))
            if caption:
                doc.add_paragraph(caption, style="Caption")


class TestResultTable:
    """Helper class to create test result tables."""

    @staticmethod
    def from_indicator_results(
        results: List[Any],
        title: str = "测试结果汇总",
    ) -> Dict[str, Any]:
        """
        Create table data from indicator results.

        Args:
            results: List of IndicatorResult objects.
            title: Table title.

        Returns:
            Dictionary with table data.
        """
        headers = ["指标名称", "计算值", "判定结果", "备注"]
        rows = []

        for result in results:
            rows.append(
                [
                    result.definition.name,
                    f"{result.calculated_value:.4f}"
                    if result.calculated_value is not None
                    else "N/A",
                    result.judgment.value,
                    result.error_message or "",
                ]
            )

        return {"title": title, "headers": headers, "rows": rows}
